import docx 
import os
import re


document = docx.Document("./3.docx")


def addFile(name):
    with open(name, 'r') as f:
        content = f.read()
        document.add_paragraph('Файл '+ os.path.basename(f.name) + ':')
        document.add_paragraph()
        new_paragraph = document.add_paragraph()
        new_paragraph.paragraph_format.first_line_indent = docx.shared.Inches(0)
        new_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
        run1 = new_paragraph.add_run(content)

        run1.font.name = 'Courier New'
        run1.font.size = docx.shared.Pt(10)
        document.add_paragraph()
def process_files_recursively(root_directory):
    for dirpath, dirnames, filenames in os.walk(root_directory):
        for filename in filenames:
            if(re.search(r"\.cc$", filename) or re.search(r"\.hh$", filename) or re.search(r"\.cpp$", filename) or re.search(r"\.h$", filename) or re.search(r"\.css$", filename) or  re.search(r"\.html$", filename) or  re.search(r"\.js$", filename)):
                print(os.path.join(dirpath, filename))
                addFile(os.path.join(dirpath, filename))

process_files_recursively(input())

document.save("res.docx")